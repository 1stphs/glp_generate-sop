import json
import os
import re
import argparse
from docx import Document
from pathlib import Path

def parse_markdown_table(md_text):
    lines = [l.strip() for l in md_text.strip().split('\n') if l.strip()]
    if len(lines) < 2: return None
    if not any('|' in l and '-' in l for l in lines): return None
    rows = []
    for line in lines:
        if re.match(r'^\|?[\-\s|]+\|?$', line): continue
        cells = [c.strip() for c in line.split('|') if c.strip() or line.count('|') > 1]
        if cells: rows.append(cells)
    return rows

def json_to_docx(input_json, output_docx):
    print(f"[*] Converting {input_json} to {output_docx}...")
    if not os.path.exists(input_json):
        print(f"Error: {input_json} not found.")
        return

    with open(input_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    doc.add_heading('试验验证报告 (模拟生成版)', 0)
    
    for item in data:
        title = item.get("section_id", "Untitled")
        content = item.get("generate_content", "")
        if not content or content.strip() == "": continue
        
        doc.add_heading(title, level=1)
        
        if "|" in content and "\n| ---" in content:
            parts = re.split(r'(\n\|.*\|.*\n\|[\s\-|]+\|.*\n(?:\|.*\|.*\n)*)', content)
            for part in parts:
                if not part.strip(): continue
                table_data = parse_markdown_table(part)
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=max(len(r) for r in table_data))
                    table.style = 'Table Grid'
                    for r_idx, row_cells in enumerate(table_data):
                        for c_idx, cell_text in enumerate(row_cells):
                            if c_idx < len(table.columns):
                                table.cell(r_idx, c_idx).text = cell_text
                else:
                    doc.add_paragraph(part.strip())
        else:
            doc.add_paragraph(content.strip())

    doc.save(output_docx)
    print(f"✨ Successfully saved Word document to {output_docx}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert generated JSON report to Word Docx")
    parser.add_argument("--input", default="report_generation/test_report_output.json", help="Input JSON path")
    parser.add_argument("--output", required=True, help="Output Docx path")
    args = parser.parse_args()
    json_to_docx(args.input, args.output)
