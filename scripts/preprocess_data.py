import os
import sys
import re
import json
import traceback
from pathlib import Path
from docx import Document
from typing import Dict, List, Any

# Add project root to sys.path to import sandbox components
sys.path.append(str(Path(__file__).parent.parent))
from sop_deeplang.sandbox.excel_parser import ExcelParser_Sandbox

class DocxParser:
    """
    Parses .docx files and extracts sections using TOC indices as ground truth.
    """
    def __init__(self, file_path: str, report_id: str = "Unknown"):
        self.file_path = Path(file_path)
        self.report_id = report_id
        if self.file_path.suffix == '.doc':
            # Note: doc to docx conversion is expected to be handled externally or via local tools
            self.file_path = self._convert_doc_to_docx(self.file_path)
        
        self.doc = Document(str(self.file_path))
        self.toc_list = self._extract_toc()
        self.sections = self._parse_sections()

    def _convert_doc_to_docx(self, doc_path: Path) -> Path:
        docx_path = doc_path.with_suffix('.docx')
        if docx_path.exists():
            return docx_path
        return doc_path 

    def _extract_toc(self) -> List[str]:
        """Extracts and cleans the Table of Contents from the document."""
        toc = []
        for p in self.doc.paragraphs:
            style_name = p.style.name.lower()
            if style_name.startswith('toc') and not style_name.startswith('toc heading'):
                text = p.text.strip()
                if text:
                    # Remove ending page numbers and dots/tabs
                    clean_text = re.sub(r'[\t\.\s]+\d+$', '', text).strip()
                    # Remove leading sequence numbers (e.g. "1.1 ", "1\t", "一、")
                    clean_text = re.sub(r'^(\d+(\.\d+)*|[\u2460-\u2473]|[一二三四五六七八九十百]+[、.])[\t\s]*', '', clean_text).strip()
                    if clean_text:
                        toc.append(clean_text)
        return toc

    def _parse_sections(self) -> Dict[str, str]:
        """Parses the document body sections based on the TOC and fixed headers."""
        sections = {}
        fixed_headers = [
            "验证报告", "GLP遵从性声明和签字页", "签字页", "质量保证声明", 
            "目录", "附表目录", "附图目录", "缩略语表", "摘要"
        ]
        
        # Combine fixed headers with TOC to form a complete list of targets
        all_targets = fixed_headers + self.toc_list
        # Use normalized (no spaces) keys for robust matching
        normalized_targets = {re.sub(r'\s+', '', t): t for t in all_targets if t}
        
        current_heading = "Header/Title"
        current_content = []

        def flush():
            nonlocal current_heading, current_content
            if current_content:
                sections[current_heading] = "\n".join(current_content).strip()
            current_content = []

        for child in self.doc.element.body.iterchildren():
            if child.tag.endswith('p'):
                paras = [p for p in self.doc.paragraphs if p._element == child]
                if not paras: continue
                para = paras[0]
                text = para.text.strip()
                if not text: continue
                
                # Skip the paragraphs that ARE the TOC entries
                if para.style.name.lower().startswith('toc'):
                    continue
                
                # Clean the body text for matching (remove leading numbers)
                body_clean = re.sub(r'^(\d+(\.\d+)*|[\u2460-\u2473]|[一二三四五六七八九十百]+[、.])[\t\s]*', '', text).strip()
                norm_text = re.sub(r'\s+', '', body_clean)
                
                is_new_section = False
                matched_heading = None
                
                if norm_text in normalized_targets:
                    is_new_section = True
                    matched_heading = normalized_targets[norm_text]
                else:
                    # Try partial match for fixed headers if they are long or contain extra text
                    for fh in fixed_headers:
                        if fh in text and len(text) < 50:
                            is_new_section = True
                            matched_heading = fh
                            break
                
                if is_new_section:
                    flush()
                    current_heading = matched_heading
                else:
                    current_content.append(text)
                    
            elif child.tag.endswith('tbl'):
                tables = [t for t in self.doc.tables if t._element == child]
                if tables:
                    md_table = self._table_to_markdown(tables[0])
                    current_content.append(md_table)

        flush()
        return sections

    def _table_to_markdown(self, table) -> str:
        rows = []
        for row in table.rows:
            rows.append([cell.text.replace('\n', '<br>').replace('|', '\\|').strip() for cell in row.cells])
        if not rows: return ""
        
        headers = rows[0]
        md = f"| {' | '.join(headers)} |\n"
        md += f"| {' | '.join(['---'] * len(headers))} |\n"
        for row in rows[1:]:
            md += f"| {' | '.join(row)} |\n"
        return md

def merge_sources(protocol_sections: Dict[str, str], report_sections: Dict[str, str], toc_order: List[str]) -> List[Dict]:
    """
    Merges Protocol and Report sections into a JSON-friendly list of dictionaries.
    toc_order comes from the report's DocxParser.toc_list.
    """
    fixed_headers = [
        "验证报告", "GLP遵从性声明和签字页", "质量保证声明", 
        "目录", "附表目录", "附图目录", "缩略语表", "摘要"
    ]
    
    # We use a mapping to align Protocol names to Report names where they differ
    HEADER_MAPPING = {
        "验证方案": "验证报告",
        "签字页": "GLP遵从性声明和签字页"
    }
    
    # Pre-process protocol sections with mapping
    mapped_protocol = {}
    for h, content in protocol_sections.items():
        standardized_h = HEADER_MAPPING.get(h, h)
        if standardized_h in mapped_protocol:
            mapped_protocol[standardized_h] += "\n\n" + content
        else:
            mapped_protocol[standardized_h] = content
            
    # Combine fixed headers and report's TOC for the final sequence
    ordered_sequence = []
    # 1. Fixed headers
    for h in fixed_headers:
        ordered_sequence.append(h)
    # 2. Add TOC items from report (avoiding duplicates with fixed headers)
    for h in toc_order:
        if h not in ordered_sequence:
            ordered_sequence.append(h)
            
    output_data = []
    for heading in ordered_sequence:
        p_content = mapped_protocol.get(heading, "")
        r_content = report_sections.get(heading, "")
        
        # If both are empty, we still keep the title but content is empty
        output_data.append({
            "section_title": heading,
            "original_content": p_content,
            "generate_content": r_content,
            "sop": ""
        })
        
        # Stop at "归档" as requested
        if heading == "归档" or heading == "9 归档" or "归档" in heading and len(heading) < 10:
             # Further ensure it is the '归档' section
             if heading.strip().endswith("归档") or heading.strip() == "归档":
                 break

    return output_data

def process_directory(dir_path: Path, output_root: Path, report_id: str = None):
    """Processes a single project directory or flat directory with a specific report_id."""
    print(f"--- Processing: {dir_path.name} (ID: {report_id or dir_path.name}) ---")
    
    # Case-insensitive doc discovery
    all_files = list(dir_path.iterdir())
    
    def matches_keywords(name, keywords):
        name_lower = name.lower()
        return any(kw.lower() in name_lower for kw in keywords)

    protocol_keywords = ["方案", "Protocol"]
    report_keywords = ["REPORT", "报告", "总结报告"]

    protocol_files = [f for f in all_files if matches_keywords(f.name, protocol_keywords) and f.suffix in ['.docx', '.doc'] and not f.name.startswith((".", ".~"))]
    report_files = [f for f in all_files if matches_keywords(f.name, report_keywords) and f.suffix in ['.docx', '.doc'] and not f.name.startswith((".", ".~"))]
    
    if not report_files:
        print(f"Skipping {dir_path.name}: No report file found.")
        return

    # If report_id is not passed, use the folder name
    final_report_id = report_id or dir_path.name
    
    # Specific logic for flat directories: filter files by the extracted report_id if possible
    current_protocol = protocol_files[0] if protocol_files else None
    current_report = report_files[0] if report_files else None
    
    # If we have multiple report files and a specific ID, filter for that ID
    if report_id:
        matching_reports = [f for f in report_files if report_id in f.name]
        if matching_reports:
             current_report = matching_reports[0]
        matching_protocols = [f for f in protocol_files if report_id in f.name]
        if matching_protocols:
             current_protocol = matching_protocols[0]

    report_out_dir = output_root / final_report_id
    report_out_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        # Report parser is primary as it provides the TOC order
        report_parser = DocxParser(str(current_report), final_report_id)
        protocol_parser = DocxParser(str(current_protocol), final_report_id) if current_protocol else None
        
        protocol_sections = protocol_parser.sections if protocol_parser else {}
        report_sections = report_parser.sections
        
        # Merge using the report's TOC sequence
        merged_json = merge_sources(protocol_sections, report_sections, report_parser.toc_list)
        
        out_file = report_out_dir / "filtered_data.json"
        with open(out_file, "w", encoding="utf-8") as f:
            json.dump(merged_json, f, ensure_ascii=False, indent=4)
        print(f"Successfully saved to {out_file}")
            
    except Exception as e:
        print(f"Error processing {final_report_id}: {e}")
        # traceback.print_exc()

def main():
    project_root = Path(__file__).parent.parent
    base_dir = project_root / "original_docx"
    output_root = project_root / "data_parsed"
    
    if not base_dir.exists():
        print(f"Source directory {base_dir} not found.")
        return

    for exp_type_dir in base_dir.iterdir():
        if not exp_type_dir.is_dir() or exp_type_dir.name.startswith("."):
            continue
            
        print(f"\n======== Processing Experiment Type: {exp_type_dir.name} ========")
        exp_output_root = output_root / exp_type_dir.name
        
        # Check if it has sub-directories (Nested Mode like BV)
        sub_dirs = [d for d in exp_type_dir.iterdir() if d.is_dir() and not d.name.startswith(".")]
        
        if sub_dirs:
            # Mode A: Nested (BV Type)
            for project_dir in sub_dirs:
                process_directory(project_dir, exp_output_root)
        else:
            # Mode B: Flat (Direct files)
            all_exp_files = list(exp_type_dir.iterdir())
            report_keywords = ["REPORT", "报告", "总结报告"]
            
            report_files = [f for f in all_exp_files if any(kw.lower() in f.name.lower() for kw in report_keywords) 
                           and f.suffix in ['.docx', '.doc'] and not f.name.startswith((".", ".~"))]
            
            processed_ids = set()
            for r_file in report_files:
                # Extract ID: Match common pattern like NS24461HL01 or SS23461IR01
                id_match = re.search(r'([A-Z]{2}\d{5}[A-Z]{2}\d{2})', r_file.name)
                if id_match:
                    rid = id_match.group(1)
                    if rid not in processed_ids:
                        process_directory(exp_type_dir, exp_output_root, report_id=rid)
                        processed_ids.add(rid)
                else:
                    # Skip templates or non-project files
                    if "模板" not in r_file.name:
                        print(f"⚠️ Could not extract ID from {r_file.name}, skipping.")

if __name__ == "__main__":
    main()
